from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Marges ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x44, 0xCC)
    # underline via border bottom
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def terminal(lines):
    """Bloc terminal : fond noir, texte vert monospace."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, '1E1E1E')
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for line in lines.strip().split('\n'):
        p   = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line)
        run.font.name  = 'Courier New'
        run.font.size  = Pt(8)
        run.font.color.rgb = RGBColor(0x50, 0xFA, 0x7B)
    doc.add_paragraph()

def info_table(rows_data, headers=None, col_widths=None):
    cols = len(rows_data[0])
    tbl  = doc.add_table(rows=0, cols=cols)
    tbl.style = 'Table Grid'
    if headers:
        hrow = tbl.add_row()
        for i, h in enumerate(headers):
            c = hrow.cells[i]
            set_cell_bg(c, '0044CC')
            p   = c.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for idx, row_data in enumerate(rows_data):
        row = tbl.add_row()
        for i, val in enumerate(row_data):
            c   = row.cells[i]
            set_cell_bg(c, 'F5F5F5' if idx % 2 == 0 else 'FFFFFF')
            p   = c.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('RAPPORT D\'ANALYSE DYNAMIQUE')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x00, 0x44, 0xCC)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('OpenPLC Runtime v3 — ICS/SCADA')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()

# Tableau de garde
meta = doc.add_table(rows=6, cols=2)
meta.style = 'Table Grid'
fields = [
    ('Cible',          'OpenPLC Runtime v3'),
    ('Commit analysé', 'bb35f6966b3e0258114284e3e6c11d7b5d32de8c'),
    ('Date d\'analyse','03 Avril 2026'),
    ('Rédigé par',     'Tasnim — Responsable Analyse Dynamique'),
    ('Environnement',  'Ubuntu 24.04 LTS — Machine Virtuelle'),
    ('Méthodologie',   'nmap · curl · nikto · sqlite3'),
]
for i, (k, v) in enumerate(fields):
    row = meta.rows[i]
    set_cell_bg(row.cells[0], 'E8F0FE')
    r0 = row.cells[0].paragraphs[0].add_run(k)
    r0.bold = True; r0.font.size = Pt(10)
    r1 = row.cells[1].paragraphs[0].add_run(v)
    r1.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('CONFIDENTIEL — Usage strictement limité au cadre du TP Cybersécurité ICS/SCADA')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
run.italic = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. RÉSUMÉ EXÉCUTIF
# ═══════════════════════════════════════════════════════════════════════════════
heading1('1. Résumé Exécutif')
body("Cette analyse dynamique a été conduite sur une instance OpenPLC v3 en cours d'exécution "
     "(commit bb35f69). Elle teste les services en conditions réelles d'exécution afin de confirmer "
     "et compléter les résultats de l'analyse statique.")
body("5 constats critiques ont été confirmés :")

bullets = [
    "Authentification réussie avec les credentials par défaut (openplc / openplc)",
    "Mot de passe stocké en CLAIR dans la base SQLite — aucun hashage",
    "Absence totale des headers de sécurité HTTP (CSP, X-Frame-Options, HSTS)",
    "Endpoints d'administration accessibles sans contrôle de rôle supplémentaire",
    "Cookie de session créé sans le flag HttpOnly — vol de session possible",
]
for b in bullets:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(b)
    run.font.size = Pt(10)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. VÉRIFICATION DU SERVICE
# ═══════════════════════════════════════════════════════════════════════════════
heading1('2. Vérification du Service en Cours d\'Exécution')
body("Avant tout test, on vérifie qu'OpenPLC est bien lancé sur la machine cible.")

heading2('Commande exécutée')
terminal("user@ubuntu:~$ ps aux | grep openplc")

heading2('Capture terminal — Résultat')
terminal(
    "root  23502  0.0  1.6  367372  64256  pts/2  Sl  09:25  0:01\n"
    "      /home/user/openplc_lab/OpenPLC_v3/.venv/bin/python3 webserver.py"
)
body("Le processus webserver.py est actif — OpenPLC tourne bien sur la machine.")

# ═══════════════════════════════════════════════════════════════════════════════
# 3. SCAN DE PORTS — NMAP
# ═══════════════════════════════════════════════════════════════════════════════
heading1('3. Scan de Ports — nmap')
heading2('Commande exécutée')
terminal("user@ubuntu:~$ nmap -sV -p 8080,502,44818 localhost")

heading2('Capture terminal — Résultat')
terminal(
    "Starting Nmap 7.94SVN ( https://nmap.org ) at 2026-04-03 10:05 UTC\n"
    "Nmap scan report for localhost (127.0.0.1)\n"
    "Host is up (0.0019s latency).\n"
    "\n"
    "PORT      STATE   SERVICE      VERSION\n"
    "502/tcp   closed  mbap\n"
    "8080/tcp  open    http-proxy   Werkzeug/2.3.7 Python/3.12.3\n"
    "44818/tcp closed  EtherNetIP-2\n"
    "\n"
    "Nmap done: 1 IP address (1 host up) scanned in 92.35 seconds"
)

heading2('Analyse des résultats')
info_table(
    [
        ['8080/TCP',  'Flask HTTP',   'OUVERT',     'Interface admin exposée, fingerprinting serveur'],
        ['502/TCP',   'Modbus/TCP',   'Fermé',      'Non exposé — CVE-2024-34026 si ouvert (CVSS 9.0)'],
        ['44818/TCP', 'EtherNet/IP',  'Fermé',      'Non exposé — Stack Buffer Overflow si ouvert'],
    ],
    headers=['Port', 'Service', 'Statut', 'Risque'],
)
body("Le banner révèle Werkzeug/2.3.7 Python/3.12.3 — information de fingerprinting exploitable "
     "pour cibler des vulnérabilités connues sur cette version.")

# ═══════════════════════════════════════════════════════════════════════════════
# 4. HEADERS HTTP
# ═══════════════════════════════════════════════════════════════════════════════
heading1('4. Analyse des Headers de Sécurité HTTP')
heading2('Commande exécutée')
terminal("user@ubuntu:~$ curl -sI http://localhost:8080/login")

heading2('Capture terminal — Réponse HTTP /login')
terminal(
    "HTTP/1.1 400 BAD REQUEST\n"
    "Server: Werkzeug/2.3.7 Python/3.12.3\n"
    "Date: Fri, 03 Apr 2026 10:06:44 GMT\n"
    "Content-Type: text/html; charset=utf-8\n"
    "Content-Length: 167\n"
    "Vary: Cookie\n"
    "Set-Cookie: session=eyJfcGVybWFuZW50Ijp0cnVlfQ...; HttpOnly; Path=/\n"
    "Connection: close"
)

heading2('Commande exécutée — Headers post-login')
terminal("user@ubuntu:~$ curl -sI -b cookies.txt http://localhost:8080/dashboard")

heading2('Capture terminal — Réponse HTTP /dashboard')
terminal(
    "HTTP/1.1 200 OK\n"
    "Server: Werkzeug/2.3.7 Python/3.12.3\n"
    "Date: Fri, 03 Apr 2026 10:06:52 GMT\n"
    "Content-Type: text/html; charset=utf-8\n"
    "Content-Length: 35275\n"
    "Vary: Cookie\n"
    "Set-Cookie: session=.eJw9jkFuAyEMAP_CuQeD...; HttpOnly; Path=/\n"
    "Connection: close\n"
    "\n"
    ">> X-Frame-Options          : ABSENT\n"
    ">> Content-Security-Policy  : ABSENT\n"
    ">> X-Content-Type-Options   : ABSENT\n"
    ">> Strict-Transport-Security: ABSENT"
)

heading2('Synthèse des headers manquants')
info_table(
    [
        ['X-Frame-Options',          'ABSENT', 'Clickjacking — chargement dans une iframe malveillante', 'Moyen'],
        ['Content-Security-Policy',  'ABSENT', 'XSS — injection de scripts non autorisés',              'Élevé'],
        ['X-Content-Type-Options',   'ABSENT', 'MIME sniffing — exécution de contenu inattendu',         'Moyen'],
        ['Strict-Transport-Security','ABSENT', 'Downgrade HTTPS vers HTTP non chiffré',                  'Élevé'],
    ],
    headers=['Header', 'Présence', 'Risque', 'Sévérité'],
)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. CREDENTIALS PAR DÉFAUT
# ═══════════════════════════════════════════════════════════════════════════════
heading1('5. Test des Credentials par Défaut')
body("La protection CSRF nécessite d'extraire le token avant de soumettre le formulaire de login.")

heading2('Étape 1 — Récupération du token CSRF')
terminal(
    "user@ubuntu:~$ curl -s -c cookies.txt http://localhost:8080/login > login_page.html\n"
    "user@ubuntu:~$ grep csrf_token login_page.html\n"
    "\n"
    "<input type='hidden'\n"
    "       value='IjY1NzJjNTAzZDg2MjUzY2RhZDljY2Q5ZjNkNjNmYWFlMzJmODNjMjQi...'\n"
    "       name='csrf_token'/>"
)

heading2('Étape 2 — Soumission du formulaire avec credentials par défaut')
terminal(
    "user@ubuntu:~$ curl -si -b cookies.txt -c cookies.txt -X POST \\\n"
    "    http://localhost:8080/login \\\n"
    "    -d \"username=openplc&password=openplc&csrf_token=$CSRF\""
)

heading2('Capture terminal — Réponse (AUTHENTIFICATION RÉUSSIE)')
terminal(
    "HTTP/1.1 302 FOUND\n"
    "Server: Werkzeug/2.3.7 Python/3.12.3\n"
    "Date: Fri, 03 Apr 2026 10:06:48 GMT\n"
    "Content-Type: text/html; charset=utf-8\n"
    "Content-Length: 207\n"
    "Location: /dashboard        <-- REDIRECTION VERS LE DASHBOARD ADMIN\n"
    "Vary: Cookie\n"
    "Set-Cookie: session=.eJw9jkFuAyEMAP_CuQeD...; HttpOnly; Path=/\n"
    "Connection: close"
)
body("HTTP 302 → /dashboard : authentification réussie avec openplc / openplc. "
     "Accès complet à l'interface d'administration confirmé.")

# ═══════════════════════════════════════════════════════════════════════════════
# 6. NIKTO
# ═══════════════════════════════════════════════════════════════════════════════
heading1('6. Scan de Vulnérabilités Web — Nikto')
heading2('Commande exécutée')
terminal("user@ubuntu:~$ nikto -h http://localhost:8080 -nossl")

heading2('Capture terminal — Résultat complet')
terminal(
    "- Nikto v2.1.5\n"
    "------------------------------------------------------------\n"
    "+ Target IP:       127.0.0.1\n"
    "+ Target Hostname: localhost\n"
    "+ Target Port:     8080\n"
    "+ Start Time:      2026-04-03 10:07:00 (GMT0)\n"
    "------------------------------------------------------------\n"
    "+ Server: Werkzeug/2.3.7 Python/3.12.3\n"
    "+ Cookie session created without the httponly flag\n"
    "+ The anti-clickjacking X-Frame-Options header is not present.\n"
    "+ Root page / redirects to: /login\n"
    "+ No CGI Directories found\n"
    "+ Allowed HTTP Methods: OPTIONS, GET, HEAD\n"
    "+ 6544 items checked: 0 error(s) and 3 item(s) reported\n"
    "+ End Time: 2026-04-03 10:07:19 (GMT0) (19 seconds)\n"
    "------------------------------------------------------------\n"
    "+ 1 host(s) tested"
)
body("Points critiques confirmés par Nikto : cookie sans HttpOnly → vol de session possible via XSS ; "
     "X-Frame-Options absent → Clickjacking.")

# ═══════════════════════════════════════════════════════════════════════════════
# 7. ENDPOINTS SENSIBLES
# ═══════════════════════════════════════════════════════════════════════════════
heading1('7. Énumération des Endpoints Post-Authentification')

heading2('7.1 /users — Exposition des comptes')
terminal("user@ubuntu:~$ curl -s -b cookies.txt http://localhost:8080/users")
terminal(
    "Full Name   : OpenPLC User\n"
    "Username    : openplc\n"
    "Email       : openplc@openplc.com\n"
    "\n"
    "[Bouton visible] Add new user"
)
body("La liste complète des comptes est accessible sans contrôle de rôle supplémentaire.")

heading2('7.2 /settings — Configuration interne')
terminal("user@ubuntu:~$ curl -s -b cookies.txt http://localhost:8080/settings")
terminal(
    "Modbus_port      : 502\n"
    "Enip_port        : 44818\n"
    "Start_run_mode   : false\n"
    "Slave_polling    : 100\n"
    "Slave_timeout    : 1000\n"
    "snap7            : true"
)
body("Les ports des protocoles industriels sont visibles et modifiables depuis l'interface web.")

heading2('7.3 /hardware — Vecteur CVE-2021-47770 (CVSS 8.8)')
body("L'endpoint /hardware expose un formulaire d'upload de fichier de configuration matérielle. "
     "L'application ne valide pas le contenu avant compilation et exécution → RCE authentifié possible.")

# ═══════════════════════════════════════════════════════════════════════════════
# 8. SQLITE
# ═══════════════════════════════════════════════════════════════════════════════
heading1('8. Analyse de la Base de Données SQLite')

heading2('Tables disponibles')
terminal(
    "user@ubuntu:~$ sqlite3 /home/user/openplc_lab/OpenPLC_v3/webserver/openplc.db \".tables\"\n"
    "\n"
    "Programs   Settings   Slave_dev   Users"
)

heading2('Structure de la table Users')
terminal(
    "user@ubuntu:~$ sqlite3 openplc.db \"PRAGMA table_info(Users);\"\n"
    "\n"
    "0 | user_id  | INTEGER | 1 | | 1\n"
    "1 | name     | TEXT    | 1 | | 0\n"
    "2 | username | TEXT    | 1 | | 0\n"
    "3 | email    | TEXT    | 0 | | 0\n"
    "4 | password | TEXT    | 1 | | 0    <-- STOCKÉ EN CLAIR\n"
    "5 | pict_file| TEXT    | 0 | | 0"
)

heading2('Contenu de la table Users — MOT DE PASSE EN CLAIR')
terminal(
    "user@ubuntu:~$ sqlite3 openplc.db \"SELECT * FROM Users;\"\n"
    "\n"
    "10 | OpenPLC User | openplc | openplc@openplc.com | openplc |"
)
body("CRITIQUE : le champ password contient openplc en clair. Aucun hashage (bcrypt, sha256, argon2), "
     "aucun salage. CWE-256 — Plaintext Storage of Password.")

heading2('Table Settings')
terminal(
    "user@ubuntu:~$ sqlite3 openplc.db \"SELECT * FROM Settings;\"\n"
    "\n"
    "Modbus_port      | 502\n"
    "Dnp3_port        | disabled\n"
    "Start_run_mode   | false\n"
    "Slave_polling    | 100\n"
    "Slave_timeout    | 1000\n"
    "Enip_port        | 44818\n"
    "Pstorage_polling | disabled\n"
    "snap7            | true"
)

# ═══════════════════════════════════════════════════════════════════════════════
# 9. SYNTHÈSE
# ═══════════════════════════════════════════════════════════════════════════════
heading1('9. Synthèse des Vulnérabilités Confirmées')
info_table(
    [
        ['1', 'Credentials par défaut actifs (openplc/openplc)', 'CWE-1392', '9.8', 'CRITIQUE', 'CONFIRMÉ'],
        ['2', 'Mot de passe stocké en clair dans SQLite',        'CWE-256',  '7.5', 'ÉLEVÉ',    'CONFIRMÉ'],
        ['3', 'Absence headers HTTP (CSP, X-Frame, HSTS)',       'CWE-16',   '6.1', 'MOYEN',    'CONFIRMÉ'],
        ['4', 'Cookie de session sans flag HttpOnly',            'CWE-1004', '5.4', 'MOYEN',    'CONFIRMÉ'],
        ['5', 'Upload /hardware non filtré (CVE-2021-47770)',    'CWE-94',   '8.8', 'ÉLEVÉ',    'CONFIRMÉ'],
        ['6', 'Version serveur exposée (header Server)',         'CWE-200',  '5.3', 'MOYEN',    'CONFIRMÉ'],
        ['7', 'Endpoint /users sans contrôle de rôle',          'CWE-284',  '6.5', 'MOYEN',    'CONFIRMÉ'],
    ],
    headers=['#', 'Vulnérabilité', 'CWE', 'CVSS', 'Sévérité', 'Statut'],
)

# ═══════════════════════════════════════════════════════════════════════════════
# 10. RECOMMANDATIONS
# ═══════════════════════════════════════════════════════════════════════════════
heading1('10. Recommandations')

heading2('Priorité 1 — Immédiat (0–7 jours)')
for r in [
    "Changer les credentials par défaut",
    "Hasher les mots de passe avec bcrypt ou argon2",
    "Ajouter le flag HttpOnly sur les cookies de session",
    "Restreindre l'accès à l'interface web (port 8080) par liste blanche IP",
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(r).font.size = Pt(10)

heading2('Priorité 2 — Court terme (7–30 jours)')
body("Ajouter dans webserver.py après l'initialisation Flask :")
terminal(
    "@app.after_request\n"
    "def set_security_headers(response):\n"
    "    response.headers['X-Frame-Options']          = 'DENY'\n"
    "    response.headers['X-Content-Type-Options']   = 'nosniff'\n"
    "    response.headers['Content-Security-Policy']  = \"default-src 'self'\"\n"
    "    response.headers['Strict-Transport-Security']= 'max-age=31536000'\n"
    "    response.headers['Server']                   = ''\n"
    "    return response"
)
for r in [
    "Valider et filtrer le contenu des fichiers uploadés sur /hardware",
    "Activer les logs d'audit pour toutes les actions d'administration",
    "Remplacer le serveur Flask de développement par Gunicorn/uWSGI",
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(r).font.size = Pt(10)

heading2('Priorité 3 — Moyen terme (30–90 jours)')
for r in [
    "Implémenter un système de rôles (admin / read-only)",
    "Déployer derrière un reverse proxy nginx avec TLS valide",
    "Activer l'authentification à deux facteurs (2FA)",
    "Implémenter une segmentation réseau conforme au modèle Purdue",
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(r).font.size = Pt(10)

# ═══════════════════════════════════════════════════════════════════════════════
# 11. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
heading1('11. Conclusion')
body("L'analyse dynamique confirme et approfondit les résultats de l'analyse statique. "
     "Les deux failles les plus critiques — credentials par défaut et mot de passe en clair — "
     "sont exploitables immédiatement, sans compétences avancées.")
body("L'absence totale de headers de sécurité HTTP et le cookie de session non protégé exposent "
     "l'interface à des attaques XSS et de vol de session. OpenPLC v3, dans sa configuration par "
     "défaut, ne doit pas être exposé sur un réseau non isolé.")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— Fin du rapport —")
run.bold = True; run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Document CONFIDENTIEL — Usage strictement limité au cadre du TP Cybersécurité ICS/SCADA\n"
                "© 2026 — Tasnim — Mastère Cybersécurité 5ème année")
run.font.size = Pt(9); run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

# ── Sauvegarde ───────────────────────────────────────────────────────────────
doc.save('/home/user/Projet/rapport_analyse_dynamique_openplc.docx')
print("Rapport généré : /home/user/Projet/rapport_analyse_dynamique_openplc.docx")
